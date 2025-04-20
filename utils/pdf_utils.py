import os
import os
import fitz  # PyMuPDF
from typing import Optional
from pdf2docx import Converter
import docx
import tempfile

def extract_text_from_pdf(pdf_path: str) -> Optional[str]:
    """
    Extract text content from a PDF file using PyMuPDF first, then fallback to pdf2docx if needed
    """
    try:
        # First validate the PDF
        if not validate_pdf(pdf_path):
            print(f"[DEBUG] Invalid PDF format: {pdf_path}")
            return None

        print(f"[DEBUG] Attempting direct PDF text extraction with PyMuPDF: {pdf_path}")
        
        try:
            # Try direct PyMuPDF extraction first
            doc = fitz.open(pdf_path)
            text_parts = []
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                print(f"[DEBUG] Processing page {page_num + 1}/{doc.page_count}")
                
                # Get text with more detailed parameters
                page_text = page.get_text(
                    "text",  # Get plain text
                    sort=True,  # Sort blocks by reading order
                    flags=fitz.TEXT_PRESERVE_LIGATURES | fitz.TEXT_PRESERVE_WHITESPACE  # Preserve formatting
                ).strip()
                
                if page_text:
                    text_parts.append(page_text)
                    print(f"[DEBUG] Page {page_num + 1}: Found {len(page_text)} characters")
                else:
                    print(f"[DEBUG] Page {page_num + 1}: No text found")
            
            doc.close()
            
            if text_parts:
                final_text = '\n\n'.join(text_parts)
                print(f"[DEBUG] Successfully extracted {len(text_parts)} pages of text")
                print(f"[DEBUG] Total text length: {len(final_text)} characters")
                print(f"[DEBUG] First 200 chars:\n{final_text[:200]}...")
                return final_text
            else:
                print("[DEBUG] No text found with PyMuPDF, trying pdf2docx fallback")
        except Exception as e:
            print(f"[DEBUG] Error in PyMuPDF extraction: {str(e)}")
            print("[DEBUG] Falling back to pdf2docx method")
        
        # Fallback to pdf2docx method
        print(f"[DEBUG] Converting PDF to DOCX: {pdf_path}")
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_docx:
            docx_path = tmp_docx.name
        
        try:
            # Convert PDF to DOCX
            cv = Converter(pdf_path)
            cv.convert(docx_path)
            cv.close()
            
            print(f"[DEBUG] PDF converted to DOCX: {docx_path}")
            
            # Extract text from DOCX
            doc = docx.Document(docx_path)
            text_parts = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())
            
            # Clean up
            os.unlink(docx_path)
            
            if text_parts:
                final_text = '\n'.join(text_parts)
                print(f"[DEBUG] Successfully extracted {len(text_parts)} text blocks via DOCX")
                print(f"[DEBUG] Total text length: {len(final_text)} characters")
                print(f"[DEBUG] First 200 chars:\n{final_text[:200]}...")
                return final_text
            else:
                print(f"[DEBUG] No text found in converted document")
                return None
                
        except Exception as e:
            print(f"[DEBUG] Error in pdf2docx conversion: {str(e)}")
            if os.path.exists(docx_path):
                os.unlink(docx_path)
            return None
            
    except Exception as e:
        print(f"[DEBUG] Error extracting text from PDF: {str(e)}")
        return None

def validate_pdf(file_path: str) -> bool:
    """
    Validate if the file is a valid PDF
    """
    try:
        print(f"[DEBUG] Starting PDF validation for: {file_path}")
        
        # Check if file exists and is not empty
        if not os.path.exists(file_path):
            print(f"[DEBUG] File does not exist: {file_path}")
            return False
            
        file_size = os.path.getsize(file_path)
        print(f"[DEBUG] File size: {file_size} bytes")
        
        if file_size == 0:
            print(f"[DEBUG] File is empty: {file_path}")
            return False
            
        # Try to read first few bytes to check PDF signature
        with open(file_path, 'rb') as f:
            header = f.read(1024)
            print(f"[DEBUG] File header (first 20 bytes): {header[:20]}")
            if not header.startswith(b'%PDF'):
                print(f"[DEBUG] Invalid PDF header: does not start with %PDF")
                return False
            
        # Try to open as PDF
        print(f"[DEBUG] Attempting to open with PyMuPDF...")
        doc = fitz.open(file_path)
        
        # Check if it's a valid PDF
        is_pdf = doc.is_pdf
        print(f"[DEBUG] PyMuPDF is_pdf check: {is_pdf}")
        
        if not is_pdf:
            print(f"[DEBUG] Not a valid PDF file according to PyMuPDF: {file_path}")
            doc.close()
            return False
            
        # Check if it has any pages
        page_count = doc.page_count
        print(f"[DEBUG] PDF page count: {page_count}")
        
        if page_count == 0:
            print(f"[DEBUG] PDF has no pages: {file_path}")
            doc.close()
            return False
            
        # Try to access first page metadata
        try:
            first_page = doc[0]
            page_size = first_page.rect
            print(f"[DEBUG] First page size: {page_size}")
        except Exception as e:
            print(f"[DEBUG] Error accessing first page: {str(e)}")
            doc.close()
            return False
            
        doc.close()
        print(f"[DEBUG] PDF validation successful")
        return True
    except Exception as e:
        print(f"[DEBUG] Error validating PDF: {str(e)}")
        return False
